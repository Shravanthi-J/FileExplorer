from tkinter import *
from tkinter import messagebox
import os, string, re
from docx import Document
import logging as log

class FileExplorerLayout(Tk):
    """ Desktop app to search file by using Patterns/by entering Exact File name
        and merge all .txt file to one single .txt file or all .docx files to one
        single .docx file"""

    #Desktop Path to save merged documents/files
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']),'Desktop')

    #Drives List
    drives = [f'{i}:\\' for i in string.ascii_uppercase if os.path.exists(f'{i}:\\')]

    searched_files_list = []
    filtered_files_list = []
    log.basicConfig(filename='FileSearchandMerger_log.txt', level=log.DEBUG, format='%(asctime)s:%(levelname)s:%(message)s',datefmt='%d/%m/%Y %H:%M:%S')


    def __init__(self):
        super().__init__() #Is-A Relationship (Extending Tkinter class)
        #configure tkinter window
        self.configure(bg='lightgrey')
        self.title('File Search and Merge')
        self.geometry("800x520")  # Screen/Window size , width=800 and height=520
        self.resizable(0, 0)  # To prevent window getting resized
        self.createFramesButtonsListbox()




    def createFramesButtonsListbox(self):
        ''' tkinter Window configurations - labels, buttons, Listbox '''

        my_frame=Frame(self)
        large_font = ('Verdana',14) #To increase font size in Entry box
        self.entry = Entry(self,fg="black", bg="white",font=large_font) #Search box
        self.entry.pack(side=TOP)
        b = Button(self, text='Search', command=self.fileSearch) #Search Button
        b.pack()
        b.place(x=600,y=0.5)
        sbv=Scrollbar(my_frame,orient='vertical')
        sbh=Scrollbar(my_frame,orient='horizontal')
        self.listbox=Listbox(my_frame,height=22,width=100,xscrollcommand=sbh.set,yscrollcommand=sbv.set,selectmode = "multiple")
        sbv.config(command=self.listbox.yview)
        sbv.pack(side=RIGHT,fill=Y)
        sbh.config(command=self.listbox.xview)
        sbh.pack(side="bottom",fill=X)
        my_frame.pack()
        self.listbox.pack(side=TOP,expand=YES,pady=15)
        b1 = Button(self, text='show .txt files', command=self.filterTxtFiles)
        b1.pack()
        b1.place(x=250,y=445)
        b2 = Button(self, text='show .docx files', command=self.filterDocxFiles)
        b2.pack()
        b2.place(x=450,y=440)
        b3 = Button(self, text='Merge only selected files',command=self.selectedItem)
        b3.pack()
        b3.place(x=250,y=480)
        b4 = Button(self, text='Merge all files',command=self.allFiles)
        b4.pack()
        b4.place(x=450,y=480)
        log.info('Tkinter Window opened')

    def fileSearch(self):
        ''' Search for files entered in search box in all the directories
            Wildcard search also included along with exact file name, Below 3 Search Patterns included
            1. * - List all Files
            2. *.ext - Lists all files with given extension
            3. A - Lists all files starting with A '''
        try:
            if self.entry.get() == '':
                messagebox.showinfo('Information','Please enter file name to search')
                log.error('File name not entered')
            else:
                log.info('Search Pattern entered :' + self.entry.get())
                for drive in FileExplorerLayout.drives:
                    for root, dirs, files in os.walk('D:\\'):
                        for file in files:
                            if file.startswith(self.entry.get()):
                                FileExplorerLayout.searched_files_list.append(root + "\\" + file)
                            elif self.entry.get().startswith('*.'):
                                ext = self.entry.get().split('*.')[1]
                                if file.endswith('.' + ext):
                                    self.searched_files_list.append(root + "\\" + file)
                            elif self.entry.get().startswith('*'):
                                FileExplorerLayout.searched_files_list.append(root + "\\" + file)

            for file in FileExplorerLayout.searched_files_list:
                self.listbox.insert(END, file)
            log.info('Searched filed list shown on Listbox')
        except Exception as e:
            print('There was an exception in Searching file:', e)
            log.exception('Search Eception',e)

    def filterTxtFiles(self):
        """Filters only files with .txt extension from Searched file list"""
        try:
            self.listbox.delete(0, END)
            for file in FileExplorerLayout.searched_files_list:
                if file.endswith('.txt'):
                    self.listbox.insert(END, file)
            log.info('Filtered txt files from search list')
        except Exception as e:
            print('There was an exception in Filtering .txt files:', e)
            log.exception('txt files not filtered',e)


    def filterDocxFiles(self):
        """Filters only files with .docx extension from Searched file list"""
        try:
            self.listbox.delete(0, END)
            for file in FileExplorerLayout.searched_files_list:
                if file.endswith('.docx'):
                    self.listbox.insert(END, file)
            log.info('Filtered docx files from search list')
        except Exception as e:
            print('There was an exception in Filtering .docx files:', e)
            log.exception('docx files not filtered', e)

    def mergeTxtAndDocFiles(self,filtered_files_list):
        """ Merges text or doc files from Searched File List"""

        merged_doc = Document()
        try:
            f1 = open(os.path.join(FileExplorerLayout.desktop,'Merged_txt_Files_challenge2.txt'), 'w')
            for file in FileExplorerLayout.filtered_files_list:
                if file.endswith('.docx'):
                    doc = Document(file)
                    merged_doc.add_heading(file.rstrip('.docx'), 1)
                    for para in doc.paragraphs:
                        text = para.text
                        merged_doc.add_paragraph(text)
                elif file.endswith('.txt'):
                    f = open(file, 'r')
                    f1.write('-----' + f.name + '-----\n')
                    f1.write(f.read())
                    f1.write('\n\n')
            merged_doc.save(os.path.join(FileExplorerLayout.desktop, 'Combined_Docx_Files_challenge2.docx'))
            messagebox.showinfo('information', '.txt or .docx files have been merged')
            log.info('Files merged')
        except Exception as e:
            print('Exception occured while merging files:', e)
            log.exception('files not merged', e)
        finally:
            f.close()
            f1.close()
            log.info('Opened txt files are closed')

    def selectedItem(self):
        " Stores Selected File to a List "
        try:
            for selectedFile in self.listbox.curselection():
                print(self.listbox.get(selectedFile))
                FileExplorerLayout.filtered_files_list.append(self.listbox.get(selectedFile))
            self.mergeTxtAndDocFiles(FileExplorerLayout.filtered_files_list)
            log.info('Stored selected files to List')
        except Exception as e:
            print('Unable to select files',e)
            log.exception('Unable to select file',e)

    def allFiles(self):
        ''' Stores all files listed in Search window to a List'''

        try:
            for file in (self.listbox.get(0, END)):
                if file not in FileExplorerLayout.filtered_files_list:
                    FileExplorerLayout.filtered_files_list.append(file)
            self.mergeTxtAndDocFiles(FileExplorerLayout.filtered_files_list)
            log.info('Files stored to List')
        except Exception as e:
            print(e)
            log.exception('files were stored to list for merging',e)

